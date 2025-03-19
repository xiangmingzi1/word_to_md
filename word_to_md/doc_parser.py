from docx import Document
from docx.oxml.ns import qn
from .config import NAMESPACES


def get_outline_level(para):
    """获取段落的大纲级别"""
    try:
        return int(para._element.xpath('.//w:outlineLvl')[0].get(qn('w:val')))
    except:
        return None


def extract_text_with_superscript_subscript(paragraph):
    """提取段落文本，保留上标和下标，转换为 LaTeX 格式"""
    text = ""
    for run in paragraph.runs:
        run_text = run.text
        if not run_text:
            continue

        is_sup = run.element.find('.//w:vertAlign[@w:val="superscript"]', NAMESPACES) is not None
        is_sub = run.element.find('.//w:vertAlign[@w:val="subscript"]', NAMESPACES) is not None

        if is_sup:
            text += f"$^{{{run_text}}}$"
        elif is_sub:
            text += f"$_{{{run_text}}}$"
        else:
            text += run_text
    return text.strip()


def table_to_html(table):
    """将 Word 表格转换为 HTML 格式"""
    html = "<table border='1'>\n"
    for row in table.rows:
        html += "  <tr>\n"
        for cell in row.cells:
            cell_text = " ".join(extract_text_with_superscript_subscript(p) for p in cell.paragraphs if p.text.strip())
            html += f"    <td>{cell_text}</td>\n"
        html += "  </tr>\n"
    html += "</table>"
    return html


def parse_word_document(doc_path):
    """读取并解析 Word 文档"""
    doc = Document(doc_path)
    sections = []
    current_section = {"titles": [], "text": ""}
    title_stack = []

    elements = list(doc.element.body)
    for elem in elements:
        if elem.tag.endswith('p'):
            para = doc.paragraphs[[p._element for p in doc.paragraphs].index(elem)]
            text = extract_text_with_superscript_subscript(para)
            if not text:
                continue

            outline_level = get_outline_level(para)
            if para.style.name.startswith("Heading") or (outline_level is not None):
                level = outline_level + 1 if outline_level is not None else int(para.style.name.split()[-1])
                title_stack = title_stack[:level - 1]
                title_stack.append(text)
                if current_section["text"]:
                    sections.append(current_section)
                current_section = {"titles": title_stack.copy(), "text": ""}
            else:
                current_section["text"] += text + "\n"

        elif elem.tag.endswith('tbl'):
            table = doc.tables[[t._element for t in doc.tables].index(elem)]
            table_html = table_to_html(table)
            current_section["text"] += "\n" + table_html + "\n"

    if current_section["text"]:
        sections.append(current_section)
    return sections